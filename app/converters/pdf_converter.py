import re
import os
from PyPDF2 import PdfReader
from docx import Document
from reportlab.lib.pagesizes import letter 
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer 
from reportlab.lib.styles import getSampleStyleSheet 
from flask import current_app

def clean_text_for_xml(text):
    if not text:
        return ""
    return re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)

class PDFConverter:
    def pdf_to_docx(self, input_path):
        try:
            reader = PdfReader(input_path)
            doc = Document()
            doc.add_heading('Documento Convertido de PDF', 0)
            
            for page_num, page in enumerate(reader.pages, 1):
                text = page.extract_text()
                if text.strip():
                    doc.add_heading(f'Página {page_num}', level=2)
                    clean_text = clean_text_for_xml(text)
                    
                    paragraphs = clean_text.split('\n')
                    for para in paragraphs:
                        if para.strip():
                            doc.add_paragraph(para.strip())
            
            filename = os.path.basename(input_path)
            name_without_ext = os.path.splitext(filename)[0]
            output_filename = f"{name_without_ext}_converted.docx"
            output_path = os.path.join(current_app.config['CONVERTED_FOLDER'], output_filename)
            doc.save(output_path)
            return output_path
        except Exception as e:
            raise Exception(f"Erro ao converter PDF para DOCX: {str(e)}")
    def docx_to_pdf(self, input_path):
        try:
            doc = Document(input_path)
            filename = os.path.basename(input_path)
            name_without_ext = os.path.splitext(filename)[0]
            output_filename = f"{name_without_ext}_converted.pdf"
            output_path = os.path.join(current_app.config['CONVERTED_FOLDER'], output_filename)
            pdf = SimpleDocTemplate(
                output_path, 
                pagesize=letter,
                rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=72
            )
            from reportlab.lib.styles import getSampleStyleSheet
            styles = getSampleStyleSheet()
            style_norm = styles['Normal']
            style_norm.alignment = 4
            story = []
            for para in doc.paragraphs:
                if para.text.strip():
                    texto_limpo = clean_text_for_xml(para.text)
                    style = styles['Heading1'] if para.style.name.startswith('Heading') else style_norm
                    p = Paragraph(texto_limpo, style)
                    story.append(p)
                    story.append(Spacer(1, 12))
            pdf.build(story)
            return output_path
        except Exception as e:
            raise Exception(f"Erro ao converter DOCX para PDF: {str(e)}")