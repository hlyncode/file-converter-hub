import markdown 
import os 
from flask import current_app
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import letter 
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer 
from html.parser import HTMLParser
from docx import Document
class HTMLToDocxParser(HTMLParser):
    def __init__(self, doc):
        super().__init__()
        self.doc = doc
        self.in_heading = False 
        self.heading_level = 0
    def handle_starttag(self, tag, attrs):
        if tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            self.in_heading = True 
            self.heading_level = int(tag[1])
        elif tag == 'p':
            pass
    def handle_endtag(self, tag):
        if tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            if self.current_text:
                self.doc.add_heading(self.current_text.strip(), level=self.heading_level)
                self.current_text = ""
            self.in_heading = False
        elif tag == 'p':
            if self.current_text:
                self.doc.add.paragraph(self.current_text.strip())
                self.current_text = ""
    def handle_data(self, data):
        self.current_text += data
class MarkdownConverter:
    def __init__(self):
        self.md = markdown.Markdown(extensions=['extra', 'nl2br'])
    def markdown_to_pdf(self, input_path):
        try:
            with open(input_path, 'r', encoding='utf-8') as f:
                md_content = f.read()
            html_content = self.md.convert(md_content)
            filename = os.path.basename(input_path)
            name_without_ext = os.path.basename(input_path)
            name_without_ext = os.path.splitext(filename)[0]
            output_filename = f"{name_without_ext}_converted.pdf"
            output_path = os.path.join(
                current_app.config['CONVERTED_FOLDER'],
                output_filename
            )
            doc = SimpleDocTemplate(output_path, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            from html import unescape 
            clean_text = unescape(html_content)
            clean_text = clean_text.replace('<p>', '\n').replace('</p>', '')
            clean_text = clean_text.replace('<br>', '\n')
            for line in clean_text.split('\n'):
                if line.strip():
                    p = Paragraph(line, styles['Normal'])
                    story.append(p)
                    story.append(Spacer(1,12))
            doc.build(story)
            return output_path
        except Exception as e:
            raise Exception(f"Erro ao converter Markdown para PDF: {str(e)}")
    def markdown_to_docx(self, input_path):
        try:
            with open(input_path, 'r', encoding='utf-8') as f:
                md_content = f.read()
            html_content = self.md.convert(md_content)
            filename = os.path.basename(input_path)
            name_without_ext = os.path.splitext(filename)[0]
            output_filename = f"{name_without_ext}_converted.docx"
            output_path = os.path.join(
                current_app.config['CONVERTED_FOLDER'],
                output_filename
            )
            doc = Document()
            parser = HTMLToDocxParser(doc)
            parser.feed(html_content)
            if parser.current_text:
                doc.add_paragraph(parser.current_text.strip())
            doc.save(output_path)
            return output_path 
        except Exception as e:
            raise Exception(f"Erro ao converter Markdown para DOCX: {str(e)}")
